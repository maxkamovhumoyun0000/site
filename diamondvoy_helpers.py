"""
Shared Diamondvoy: triggers, Gemini classifier/answer, progressive text edit, bot statistics.
Used by student_bot and admin_bot (no imports from student_bot).
"""
from __future__ import annotations

import asyncio
import html as html_module
import re
from typing import Literal

import aiohttp
from aiogram import Bot

from config import ALL_ADMIN_IDS, SUBJECTS
from db import count_available_daily_tests_global, get_all_users, get_dcoins, get_dpoints
from i18n import t
from ai_generator import _xai_generate_text, _xai_generate_text_stream

# --- Triggers & query extraction ---

# "Diamondvoy", common typos (e.g. Daimondvoy), optional "salom" prefix, optional space in "diamond voy".
_DIAMONDVOY_MARKER_RE = re.compile(
    r"(?i)(?:salom\s+)?(?:diamond\s*voy|daimondvoy|diamonvoy|dimondvoy|diamondvoy)"
)


def _strip_leading_md_noise(text: str) -> str:
    raw = (text or "").strip()
    low = raw.lower()
    while low.startswith("*"):
        raw = raw.lstrip("*").strip()
        low = raw.lower()
    return raw


def is_diamondvoy_chat_trigger(text: str | None) -> bool:
    if not text:
        return False
    raw = _strip_leading_md_noise(text)
    return _DIAMONDVOY_MARKER_RE.search(raw) is not None


def extract_diamondvoy_query(text: str) -> str:
    """Legacy alias: same as extract_diamondvoy_query_anywhere."""
    return extract_diamondvoy_query_anywhere(text)


def extract_diamondvoy_query_anywhere(text: str) -> str:
    """Query after «Diamondvoy» / typos / «diamond voy», anywhere in the message."""
    raw = _strip_leading_md_noise(text or "")
    m = _DIAMONDVOY_MARKER_RE.search(raw)
    if not m:
        return ""
    rest = raw[m.end() :].strip()
    return rest.strip(" ,.!?-:")


def sanitize_diamondvoy_reply(text: str) -> str:
    """Strip Markdown noise so Telegram plain-text replies stay readable."""
    if not text:
        return text
    s = text
    s = re.sub(r"(?m)^#{1,6}\s*", "", s)
    s = re.sub(r"\*{2,}", "", s)
    s = re.sub(r"_{2,}", "", s)
    s = re.sub(r" +([.,;:!?])", r"\1", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()


def detect_query_language(text: str, fallback: str = "uz") -> str:
    """
    Detect query language for Diamondvoy replies: uz/ru/en.
    Keeps logic lightweight and deterministic.
    """
    raw = (text or "").strip()
    if not raw:
        fb = ((fallback or "uz").lower())[:2]
        return fb if fb in ("uz", "ru", "en") else "uz"

    low = raw.lower()
    cyr = len(re.findall(r"[а-яё]", low, flags=re.I))
    latin = len(re.findall(r"[a-z]", low))

    # Strong Cyrillic signal -> Russian.
    if cyr >= 3 and cyr >= latin:
        return "ru"

    uz_markers = (
        "yo'q",
        "yoq",
        "bo'lsa",
        "bolsa",
        "uchun",
        "qanday",
        "nega",
        "nima",
        "qayer",
        "dars",
        "savol",
    )
    en_markers = (
        "what",
        "why",
        "how",
        "when",
        "where",
        "explain",
        "example",
        "please",
        "lesson",
        "grammar",
    )

    uz_score = sum(1 for m in uz_markers if m in low)
    en_score = sum(1 for m in en_markers if m in low)
    if uz_score > en_score:
        return "uz"
    if en_score > uz_score:
        return "en"

    if latin > 0:
        # Uzbek default in this project for latin-script ambiguity.
        return "uz"

    fb = ((fallback or "uz").lower())[:2]
    return fb if fb in ("uz", "ru", "en") else "uz"


def resolve_query_subject(question: str, subjects: list[str]) -> str:
    """
    Pick the most relevant subject from student's subjects for charging/logging.
    """
    subs = [str(s).strip().title() for s in (subjects or []) if str(s).strip()]
    if not subs:
        return "English"
    if len(subs) == 1:
        return subs[0]

    ql = (question or "").lower()
    has_ru_signal = bool(
        re.search(r"[а-яё]", ql, flags=re.I)
        or re.search(r"\b(russian|rus|рус|русский)\b", ql, flags=re.I)
    )
    has_en_signal = bool(
        re.search(
            r"\b(english|ingliz|eng|grammar|tense|vocabulary|word|present|past|future|"
            r"simple|continuous|perfect|article|preposition|sentence)\b",
            ql,
            flags=re.I,
        )
    )

    if has_ru_signal:
        for s in subs:
            if s.lower() == "russian":
                return s
    if has_en_signal:
        for s in subs:
            if s.lower() == "english":
                return s
    return subs[0]


# --- AI (Grok via ai_generator) ---


async def diamondvoy_is_subject_related(question: str, subjects: list[str]) -> bool:
    q = (question or "").strip()
    ql = q.lower()
    if not q:
        return True

    compact = re.sub(r"\s+", " ", ql).strip(" \t\r\n!?.,:;")
    conversational_allowlist = {
        "salom",
        "assalomu alaykum",
        "va alaykum assalom",
        "alaykum",
        "hello",
        "hi",
        "hey",
        "rahmat",
        "thanks",
        "thank you",
        "ok",
        "okay",
        "xo'p",
        "hop",
        "mayli",
    }
    if compact in conversational_allowlist:
        return True

    greeting_or_smalltalk = re.search(
        r"\b(salom|assalomu alaykum|alaykum|hello|hi|hey|good morning|good afternoon|good evening|"
        r"привет|здравствуйте|добрый)\b",
        ql,
    )
    conversational_flow = re.search(
        r"\b(rahmat|thanks|thank you|ok|xo'p|hop|ha|yo'q|yes|no|mayli|понял|ясно|okay|"
        r"yana|again|davom|continue|tushuntir|explain|misol|example|aniqla|clarify|"
        r"what do you mean|can you|could you|qanday|nima degani|why|nega|how)\b",
        ql,
    )
    if greeting_or_smalltalk or conversational_flow:
        return True

    abusive_or_harmful = re.search(
        r"\b(stupid|idiot|hate|kill|suicide|terror|bomb|porn|sex|nsfw|fuck|shit|"
        r"убить|суицид|террор|порно|еб|нах|"
        r"ahmoq|o'ldir|oldir|porn|jinsiy)\b",
        ql,
    )
    if abusive_or_harmful:
        return False

    subject_keywords = re.search(
        r"\b(english|ingliz|russian|rus|grammar|tense|vocabulary|word|reading|writing|speaking|"
        r"listening|translation|tarjima|grammatika|лексика|грамматика|present|past|future|"
        r"simple|continuous|perfect|article|preposition|sentence|matematika|math|tarix|history|"
        r"ona tili|arab tili|arab|fizika|physics|kimyo|chemistry|misol|masala|yechim|formula)\b",
        ql,
    )
    if subject_keywords:
        return True

    obviously_unrelated = re.search(
        r"\b(crypto|bitcoin|forex|betting|casino|lottery|manga|anime|movie|film|football|soccer|nba|"
        r"politics|president|celebrity|gossip|shopping|coupon)\b",
        ql,
    )
    if obviously_unrelated:
        return False

    # Keep normal conversational flow open for short student follow-ups.
    if len(compact.split()) <= 4:
        return True

    subject_line = ", ".join(subjects) if subjects else "English"
    prompt = (
        "You are an intent classifier for a tutoring assistant.\n"
        f"Student subjects: {subject_line}\n"
        f"Question: {question}\n"
        "Return ONLY one token: YES or NO.\n"
        "Answer YES when the message is any of the following:\n"
        "1) subject-related learning question\n"
        "2) greeting/small talk\n"
        "3) clarification/follow-up in normal chat flow\n"
        "Answer NO only when message is clearly unrelated to learning context or abusive."
    )
    try:
        async with aiohttp.ClientSession() as session:
            txt = (await _xai_generate_text(prompt, session=session)).strip().upper()
        return txt.startswith("YES")
    except Exception:
        return any(s.lower() in ql for s in subjects) or bool(
            re.search(r"\b(grammar|word|tense|reading|russian|english|salom|hello|hi|matematika|tarix|arab|ona tili)\b", ql)
        )


async def diamondvoy_gemini_answer(
    question: str,
    subjects: list[str],
    lang: str = "uz",
    *,
    is_admin_context: bool = False,
    conversation: list[dict[str, str]] | None = None,
) -> str:
    """
    Backward-compatible name: now powered by Grok/xAI.
    Uses proper separation of static instructions for Prompt Caching.
    """
    subject_line = ", ".join(subjects) if subjects else "English"
    lc = ((lang or "uz").lower())[:2]
    if lc not in ("ru", "en", "uz"):
        lc = "uz"

    # Check direct admin app version commands/queries
    if is_admin_context:
        ver_action = try_diamondvoy_app_version_action(question, is_admin=True, lang=lang)
        if ver_action is not None:
            return ver_action

    # Static System Prompts (easily cached)
    if is_admin_context:
        if lc == "ru":
            sys_prompt = (
                "Ты административный помощник образовательного бота Diamond Education.\n"
                "Администраторы спрашивают об учениках, оплате, паролях, отчётах и правилах центра — это в рамках твоей задачи.\n"
                "Не отказывай в помощи из‑за формулировки «личные данные», если речь об администрировании учебного центра.\n"
                "Слова вроде «manga» в тексте часто опечатка или шум — не уходи в тематику комиксов, трактуй запрос в контексте админки.\n"
                "Отвечай кратко, по делу, строго на русском языке (не переключайся на узбекский или английский).\n"
                "Если вопрос похож на уже заданный ранее в этом чате, никогда не начинай с фраз «мы это уже обсуждали», «как я говорил ранее» и подобных — отвечай по существу каждый раз заново, как на новый вопрос.\n"
                "Не используй Markdown-звёздочки (*, **) и заголовки #; короткие абзацы, при необходимости списки через • или цифры; в каждом блоке можно 1–2 эмодзи (📌 💡). Меняй структуру абзацев и выбор эмодзи от ответа к ответу — не повторяй один и тот же шаблон оформления."
            )
        elif lc == "en":
            sys_prompt = (
                "You are an administrative assistant for the Diamond Education (language tutoring) Telegram bot.\n"
                "Admins ask about students, enrollment, payments, passwords, schedules, and center policy — these are in scope.\n"
                "Do not refuse normal admin questions as 'personal data' or 'out of scope' when they clearly relate to running the school.\n"
                "The word 'manga' in messages is often noise or a typo; do not pivot to comics or entertainment unless the user clearly asks for that.\n"
                "If they ask to find a student or profile, explain that the bot has admin search flows and keep answers practical.\n"
                "Answer briefly in English only.\n"
                "If a question resembles one already asked earlier in this chat, never open with phrases like 'as discussed before', 'as I mentioned earlier', or similar — always answer it substantively as if it were new.\n"
                "Do not use Markdown asterisks or # headings; short paragraphs, optional • or numbered lists; 1–2 emojis per section (e.g. 📌 💡) are fine. Vary the paragraph structure and emoji choice between answers instead of repeating the same layout every time."
            )
        else:
            sys_prompt = (
                "Sen Diamond Education botining administrativ yordamchisisan.\n"
                "Adminlar talablari: o‘quvchilar, to‘lov, parollar, darslar va markaz qoidalari — bularning barchasi sening vazifang.\n"
                "«Shaxsiy ma'lumot» deb bekor qilma, agar savol o‘quv markazini boshqarishga tegishli bo‘lsa.\n"
                "Matndagi «manga» so‘zi ko‘pincha xato yoki ortiqcha — komiks mavzusiga o‘tma, so‘rovni admin kontekstida tushun.\n"
                "Javoblarni qisqa, aniq va faqat o'zbek tilida ber (boshqa tilga o'tma).\n"
                "Agar savol shu suhbatda avval so'ralgan savolga o'xshab qolsa, hech qachon «buni oldin ko'rib chiqqanmiz», «yuqorida aytdim» kabi iboralar bilan boshlama — har safar savolni yangidan, mazmunli javob ber.\n"
                "Markdown yulduzcha (*, **, ***) va # sarlavhalarsiz yoz. Qisqa bo‘limlar, kerak bo‘lsa • yoki raqamli ro‘yxat; har bir bo‘limda 1–2 emoji (masalan 📌, 💡). Har javobda paragraf tuzilishi va emojilarni almashtirib tur, bitta shablonni takrorlama."
            )
    else:
        if lc == "ru":
            sys_prompt = (
                "Ты помощник Diamondvoy для ученика.\n"
                "Отвечай по теме урока, кратко и понятно, строго на русском языке.\n"
                "Если вопрос похож на уже заданный ранее в этом чате, никогда не начинай с фраз «мы это уже обсуждали», «как я говорил ранее» и подобных — отвечай по существу каждый раз заново, как на новый вопрос.\n"
                "Не используй Markdown-звёздочки (*, **) и заголовки #; короткие абзацы, при необходимости списки через • или цифры; в каждом блоке можно 1–2 эмодзи (📌 💡). Меняй структуру абзацев и выбор эмодзи от ответа к ответу — не повторяй один и тот же шаблон оформления."
            )
        elif lc == "en":
            sys_prompt = (
                "You are Diamondvoy, a tutoring assistant.\n"
                "Answer only about lesson-related topics, briefly and clearly, in English.\n"
                "If a question resembles one already asked earlier in this chat, never open with phrases like 'as discussed before', 'as I mentioned earlier', or similar — always answer it substantively as if it were new.\n"
                "Do not use Markdown asterisks or # headings; short paragraphs, optional • or numbered lists; 1–2 emojis per section (e.g. 📌 💡) are fine. Vary the paragraph structure and emoji choice between answers instead of repeating the same layout every time."
            )
        else:
            sys_prompt = (
                "Sen Diamondvoy ismli o'quvchi yordamchi botsan.\n"
                "Faqat darsga oid, qisqa va tushunarli javob ber, o'zbek tilida.\n"
                "Agar savol shu suhbatda avval so'ralgan savolga o'xshab qolsa, hech qachon «buni oldin ko'rib chiqqanmiz», «yuqorida aytdim» kabi iboralar bilan boshlama — har safar savolni yangidan, mazmunli javob ber.\n"
                "Markdown yulduzcha (*, **, ***) va # sarlavhalarsiz yoz. Qisqa bo‘limlar, kerak bo‘lsa • yoki raqamli ro‘yxat; har bir bo‘limda 1–2 emoji (masalan 📌, 💡). Har javobda paragraf tuzilishi va emojilarni almashtirib tur, bitta shablonni takrorlama."
            )

    history_lines: list[str] = []
    for row in (conversation or [])[-12:]:
        role = str((row or {}).get("role") or "").strip().lower()
        content = str((row or {}).get("content") or "").strip()
        if role not in {"user", "assistant"} or not content:
            continue
        role_label = "USER" if role == "user" else "ASSISTANT"
        history_lines.append(f"{role_label}: {content[:1200]}")
    history_block = "\n".join(history_lines)
    context_hint = f"\n\nRecent conversation history:\n{history_block}\n" if history_block else ""

    user_prompt = f"Subject context: {subject_line}\n{context_hint}\nNew Question: {question}"

    try:
        async with aiohttp.ClientSession() as session:
            text = await _xai_generate_text(
                user_prompt, 
                session=session, 
                system_content=sys_prompt, 
                temperature=0.75  # Increased temperature to avoid repetitive identical answers
            )
        return (text or "").strip() or t(lang, "diamondvoy_answer_empty")
    except Exception:
        return t(lang, "diamondvoy_generation_error")

async def diamondvoy_generate_chat_title(text: str, lang: str = "uz") -> str:
    """Generate a short 4-5 word title for a chat using AI."""
    clean_text = (text or "").strip()[:500]
    if not clean_text:
        return "New chat"
        
    lc = ((lang or "uz").lower())[:2]
    if lc == "ru":
        prompt = f"Напиши очень короткий заголовок (максимум 4-5 слов) для этого сообщения. Без кавычек и лишних слов:\n{clean_text}"
    elif lc == "en":
        prompt = f"Write a very short title (max 4-5 words) for this message. No quotes, no extra words:\n{clean_text}"
    else:
        prompt = f"Bu xabar uchun juda qisqa (max 4-5 so'z) mavzu yozib ber. Hech qanday qo'shimcha so'z va qo'shtirnoqsiz:\n{clean_text}"
        
    try:
        async with aiohttp.ClientSession() as session:
            res = await _xai_generate_text(prompt, session=session)
        res = res.strip().strip('"').strip("'")
        if len(res) > 60:
            return res[:57] + "..."
        return res or clean_text[:40]
    except Exception:
        return clean_text[:40] + "..."


async def diamondvoy_gemini_answer_stream(
    question: str,
    subjects: list[str],
    lang: str = "uz",
    *,
    is_admin_context: bool = False,
    conversation: list[dict[str, str]] | None = None,
    system_prompt_override: str | None = None,
    temperature: float = 0.75,
):
    """
    Actual streaming output of the above logic via Grok/xAI.
    """
    subject_line = ", ".join(subjects) if subjects else "English"
    lc = ((lang or "uz").lower())[:2]
    if lc not in ("ru", "en", "uz"):
        lc = "uz"

    # Check direct admin app version commands/queries
    if is_admin_context:
        ver_action = try_diamondvoy_app_version_action(question, is_admin=True, lang=lang)
        if ver_action is not None:
            yield ver_action
            return

    # Static System Prompts (easily cached)
    if is_admin_context:
        if lc == "ru":
            sys_prompt = (
                "Ты административный помощник образовательного бота Diamond Education.\n"
                "Администраторы спрашивают об учениках, оплате, паролях, отчётах и правилах центра — это в рамках твоей задачи.\n"
                "Не отказывай в помощи из‑за формулировки «личные данные», если речь об администрировании учебного центра.\n"
                "Слова вроде «manga» в тексте часто опечатка или шум — не уходи в тематику комиксов, трактуй запрос в контексте админки.\n"
                "Отвечай кратко, по делу, строго на русском языке (не переключайся на узбекский или английский).\n"
                "Если вопрос похож на уже заданный ранее в этом чате, никогда не начинай с фраз «мы это уже обсуждали», «как я говорил ранее» и подобных — отвечай по существу каждый раз заново, как на новый вопрос.\n"
                "Не используй Markdown-звёздочки (*, **) и заголовки #; короткие абзацы, при необходимости списки через • или цифры; в каждом блоке можно 1–2 эмодзи (📌 💡). Меняй структуру абзацев и выбор эмодзи от ответа к ответу — не повторяй один и тот же шаблон оформления."
            )
        elif lc == "en":
            sys_prompt = (
                "You are an administrative assistant for the Diamond Education (language tutoring) Telegram bot.\n"
                "Admins ask about students, enrollment, payments, passwords, schedules, and center policy — these are in scope.\n"
                "Do not refuse normal admin questions as 'personal data' or 'out of scope' when they clearly relate to running the school.\n"
                "The word 'manga' in messages is often noise or a typo; do not pivot to comics or entertainment unless the user clearly asks for that.\n"
                "If they ask to find a student or profile, explain that the bot has admin search flows and keep answers practical.\n"
                "Answer briefly in English only.\n"
                "If a question resembles one already asked earlier in this chat, never open with phrases like 'as discussed before', 'as I mentioned earlier', or similar — always answer it substantively as if it were new.\n"
                "Do not use Markdown asterisks or # headings; short paragraphs, optional • or numbered lists; 1–2 emojis per section (e.g. 📌 💡) are fine. Vary the paragraph structure and emoji choice between answers instead of repeating the same layout every time."
            )
        else:
            sys_prompt = (
                "Sen Diamond Education botining administrativ yordamchisisan.\n"
                "Adminlar talablari: o‘quvchilar, to‘lov, parollar, darslar va markaz qoidalari — bularning barchasi sening vazifang.\n"
                "«Shaxsiy ma'lumot» deb bekor qilma, agar savol o‘quv markazini boshqarishga tegishli bo‘lsa.\n"
                "Matndagi «manga» so‘zi ko‘pincha xato yoki ortiqcha — komiks mavzusiga o‘tma, so‘rovni admin kontekstida tushun.\n"
                "Javoblarni qisqa, aniq va faqat o'zbek tilida ber (boshqa tilga o'tma).\n"
                "Agar savol shu suhbatda avval so'ralgan savolga o'xshab qolsa, hech qachon «buni oldin ko'rib chiqqanmiz», «yuqorida aytdim» kabi iboralar bilan boshlama — har safar savolni yangidan, mazmunli javob ber.\n"
                "Markdown yulduzcha (*, **, ***) va # sarlavhalarsiz yoz. Qisqa bo‘limlar, kerak bo‘lsa • yoki raqamli ro‘yxat; har bir bo‘limda 1–2 emoji (masalan 📌, 💡). Har javobda paragraf tuzilishi va emojilarni almashtirib tur, bitta shablonni takrorlama."
            )
    else:
        if lc == "ru":
            sys_prompt = (
                "Ты помощник Diamondvoy для ученика.\n"
                "Отвечай по теме урока, кратко и понятно, строго на русском языке.\n"
                "Если вопрос похож на уже заданный ранее в этом чате, никогда не начинай с фраз «мы это уже обсуждали», «как я говорил ранее» и подобных — отвечай по существу каждый раз заново, как на новый вопрос.\n"
                "Не используй Markdown-звёздочки (*, **) и заголовки #; короткие абзацы, при необходимости списки через • или цифры; в каждом блоке можно 1–2 эмодзи (📌 💡). Меняй структуру абзацев и выбор эмодзи от ответа к ответу — не повторяй один и тот же шаблон оформления."
            )
        elif lc == "en":
            sys_prompt = (
                "You are Diamondvoy, a tutoring assistant.\n"
                "Answer only about lesson-related topics, briefly and clearly, in English.\n"
                "If a question resembles one already asked earlier in this chat, never open with phrases like 'as discussed before', 'as I mentioned earlier', or similar — always answer it substantively as if it were new.\n"
                "Do not use Markdown asterisks or # headings; short paragraphs, optional • or numbered lists; 1–2 emojis per section (e.g. 📌 💡) are fine. Vary the paragraph structure and emoji choice between answers instead of repeating the same layout every time."
            )
        else:
            sys_prompt = (
                "Sen Diamondvoy ismli o'quvchi yordamchi botsan.\n"
                "Faqat darsga oid, qisqa va tushunarli javob ber, o'zbek tilida.\n"
                "Agar savol shu suhbatda avval so'ralgan savolga o'xshab qolsa, hech qachon «buni oldin ko'rib chiqqanmiz», «yuqorida aytdim» kabi iboralar bilan boshlama — har safar savolni yangidan, mazmunli javob ber.\n"
                "Markdown yulduzcha (*, **, ***) va # sarlavhalarsiz yoz. Qisqa bo‘limlar, kerak bo‘lsa • yoki raqamli ro‘yxat; har bir bo‘limda 1–2 emoji (masalan 📌, 💡). Har javobda paragraf tuzilishi va emojilarni almashtirib tur, bitta shablonni takrorlama."
            )

    # Reuse the exact Diamondvoy/Grok generation path for tightly-scoped
    # server jobs (for example homework grading) while allowing those jobs to
    # supply a deterministic machine-readable instruction.
    if system_prompt_override:
        sys_prompt = str(system_prompt_override)

    history_lines: list[str] = []
    for row in (conversation or [])[-12:]:
        role = str((row or {}).get("role") or "").strip().lower()
        content = str((row or {}).get("content") or "").strip()
        if role not in {"user", "assistant"} or not content:
            continue
        role_label = "USER" if role == "user" else "ASSISTANT"
        history_lines.append(f"{role_label}: {content[:1200]}")
    history_block = "\n".join(history_lines)
    context_hint = f"\n\nRecent conversation history:\n{history_block}\n" if history_block else ""

    user_prompt = f"Subject context: {subject_line}\n{context_hint}\nNew Question: {question}"

    try:
        async with aiohttp.ClientSession() as session:
            async for chunk in _xai_generate_text_stream(
                user_prompt, 
                session=session, 
                system_content=sys_prompt, 
                temperature=temperature
            ):
                yield chunk
    except Exception:
        yield t(lang, "diamondvoy_generation_error")


# --- Streaming one message ---


async def stream_diamondvoy_text_reply(
    bot: Bot,
    chat_id: int,
    text: str,
    lang: str = "uz",
    *,
    message_id: int | None = None,
) -> None:
    clean = sanitize_diamondvoy_reply((text or "").strip()) or t(lang, "diamondvoy_answer_empty")
    if len(clean) > 4096:
        clean = clean[:4093] + "..."
    prefix = t(lang, "diamondvoy_stream_prefix")
    if message_id is None:
        msg = await bot.send_message(chat_id, prefix)
        message_id = msg.message_id
    else:
        try:
            await bot.edit_message_text(f"{prefix}\n\n", chat_id=chat_id, message_id=message_id)
        except Exception:
            pass
    step = 140
    for i in range(step, len(clean) + step, step):
        chunk = clean[:i]
        try:
            await bot.edit_message_text(
                f"{prefix}\n\n{chunk}",
                chat_id=chat_id,
                message_id=message_id,
            )
        except Exception:
            pass
        await asyncio.sleep(0.12)


# --- Bot statistics (Diamondvoy) ---

_STATS_ANY = re.compile(
    r"statistika|статистик|statistics|\bstats\b|\bstatus\b|holat|статус|"
    r"nechta|сколько|how many|"
    r"\bjami\b|umumiy|умумий|всего|\btotal\b|foydalanuvchi|users|userlar|"
    r"talaba|student|o'quvchi|oquvchi|o‘quvchi|"
    r"daily test|kunlik test|test zaxira|zaxira|stock|склад|"
    r"d'coin|dcoin|рейтинг|reyting|leaderboard|"
    r"onlayn|online|онлайн|faol",
    re.I,
)

_GLOBAL_HINTS = re.compile(
    r"\bjami\b|umumiy|умумий|всего|\btotal\b|all users|barcha|hamma|butun|"
    r"nechta talaba|how many students|studentlar soni|сколько студентов|сколько пользователей|"
    r"test zaxira|daily test|zaxira|stock|global|globaln",
    re.I,
)

_PERSONAL_HINTS = re.compile(
    r"maniki|mening|мой\b|\bmy\b|faqat men|только я|meniki|"
    r"meniki|dcoinim|d'coinim|balansim|reytingim|мой баланс|my balance|my dcoin",
    re.I,
)


def try_diamondvoy_app_version_action(
    query: str,
    *,
    is_admin: bool = False,
    lang: str = "uz",
) -> str | None:
    """
    Handles version queries and automatic version update requests for Diamondvoy AI.
    Admin can ask about versions or request version updates (e.g. "student app versiyasini 1.2.0 qil").
    """
    if not is_admin:
        return None
    raw = (query or "").strip()
    if not raw:
        return None
    ql = raw.lower()

    # Must contain version-related keywords
    version_keywords = re.search(
        r"\b(versiya|versiyalar|versiyasi|version|versions|build|force update|force_update|app versiya|app version|min_version|min_build)\b",
        ql,
    )
    if not version_keywords:
        return None

    from db import get_app_version_settings, update_app_version_settings

    # Check if this is an UPDATE command
    is_set_command = any(
        k in ql
        for k in (
            "qil",
            "o'zgartir",
            "ozgartir",
            "o‘zgartir",
            "set",
            "update",
            "oshir",
            "tengla",
            "qo'y",
            "qoy",
            "o'zgar",
            "измени",
            "обнови",
            "поставить",
            "сделать",
        )
    )

    ver_match = re.search(r"\b(\d+\.\d+(?:\.\d+)?)\b", ql)
    build_match = re.search(r"\b(?:build|b)[\s:=]*(\d+)\b", ql)

    if is_set_command and ver_match:
        new_ver = ver_match.group(1)
        new_build = int(build_match.group(1)) if build_match else None

        target_student = any(k in ql for k in ("student", "o'quvchi", "oquvchi", "talaba", "ученик", "студент"))
        target_teacher = any(k in ql for k in ("teacher", "o'qituvchi", "oqituvchi", "ustoz", "учитель", "преподаватель"))

        if not target_student and not target_teacher:
            target_student = True
            target_teacher = True

        settings = get_app_version_settings()
        upd = {}

        if target_student:
            b_num = new_build if new_build is not None else settings.get("min_student_build", 1)
            upd["min_student_version"] = new_ver
            upd["min_student_build"] = b_num
        if target_teacher:
            b_num = new_build if new_build is not None else settings.get("min_teacher_build", 1)
            upd["min_teacher_version"] = new_ver
            upd["min_teacher_build"] = b_num

        if upd:
            update_app_version_settings(upd)

        updated = get_app_version_settings()

        if lang == "ru":
            return (
                "🚀 <b>Версия мобильного приложения обновлена!</b>\n\n"
                f"📱 <b>Diamond Students App</b>:\n"
                f"• Мин. версия: <code>{updated.get('min_student_version')}</code>\n"
                f"• Мин. build: <code>{updated.get('min_student_build')}</code>\n\n"
                f"👨‍🏫 <b>Diamond Teachers App</b>:\n"
                f"• Мин. версия: <code>{updated.get('min_teacher_version')}</code>\n"
                f"• Мин. build: <code>{updated.get('min_teacher_build')}</code>\n\n"
                "📌 Все пользователи с версией ниже указанной будут перенаправлены на экран обязательного обновления."
            )
        elif lang == "en":
            return (
                "🚀 <b>Mobile App Version Updated!</b>\n\n"
                f"📱 <b>Diamond Students App</b>:\n"
                f"• Min Version: <code>{updated.get('min_student_version')}</code>\n"
                f"• Min Build: <code>{updated.get('min_student_build')}</code>\n\n"
                f"👨‍🏫 <b>Diamond Teachers App</b>:\n"
                f"• Min Version: <code>{updated.get('min_teacher_version')}</code>\n"
                f"• Min Build: <code>{updated.get('min_teacher_build')}</code>\n\n"
                "📌 All users with an older app version will now see the forced update screen."
            )
        else:
            return (
                "🚀 <b>Mobil ilovalar versiyasi muvaffaqiyatli yangilandi!</b>\n\n"
                f"📱 <b>Diamond Students App</b>:\n"
                f"• Minimal versiya: <code>{updated.get('min_student_version')}</code>\n"
                f"• Minimal build: <code>{updated.get('min_student_build')}</code>\n\n"
                f"👨‍🏫 <b>Diamond Teachers App</b>:\n"
                f"• Minimal versiya: <code>{updated.get('min_teacher_version')}</code>\n"
                f"• Minimal build: <code>{updated.get('min_teacher_build')}</code>\n\n"
                "📌 Endi ushbu versiyadan past bo'lgan barcha ilovalarda majburiy yangilanish ekrani chiqadi."
            )

    current = get_app_version_settings()
    if lang == "ru":
        return (
            "🚀 <b>Текущие настройки версий мобильных приложений</b>:\n\n"
            f"📱 <b>Diamond Students App</b>:\n"
            f"• Мин. версия: <code>{current.get('min_student_version')}</code>\n"
            f"• Мин. build: <code>{current.get('min_student_build')}</code>\n"
            f"• Play Store: {current.get('student_play_store_url')}\n"
            f"• App Store: {current.get('student_app_store_url')}\n\n"
            f"👨‍🏫 <b>Diamond Teachers App</b>:\n"
            f"• Мин. версия: <code>{current.get('min_teacher_version')}</code>\n"
            f"• Мин. build: <code>{current.get('min_teacher_build')}</code>\n"
            f"• Play Store: {current.get('teacher_play_store_url')}\n"
            f"• App Store: {current.get('teacher_app_store_url')}\n\n"
            "💡 <i>Чтобы обновить версию, напишите например: «Student app версию 1.2.0 build 15 сделать»</i>"
        )
    elif lang == "en":
        return (
            "🚀 <b>Current Mobile App Version Settings</b>:\n\n"
            f"📱 <b>Diamond Students App</b>:\n"
            f"• Min Version: <code>{current.get('min_student_version')}</code>\n"
            f"• Min Build: <code>{current.get('min_student_build')}</code>\n"
            f"• Play Store: {current.get('student_play_store_url')}\n"
            f"• App Store: {current.get('student_app_store_url')}\n\n"
            f"👨‍🏫 <b>Diamond Teachers App</b>:\n"
            f"• Min Version: <code>{current.get('min_teacher_version')}</code>\n"
            f"• Min Build: <code>{current.get('min_teacher_build')}</code>\n"
            f"• Play Store: {current.get('teacher_play_store_url')}\n"
            f"• App Store: {current.get('teacher_app_store_url')}\n\n"
            "💡 <i>To update version, write e.g.: 'Set Student app version to 1.2.0 build 15'</i>"
        )
    else:
        return (
            "🚀 <b>Hozirgi Mobil Ilovalar Versiyalari Sozlamalari</b>:\n\n"
            f"📱 <b>Diamond Students App</b>:\n"
            f"• Minimal versiya: <code>{current.get('min_student_version')}</code>\n"
            f"• Minimal build: <code>{current.get('min_student_build')}</code>\n"
            f"• Google Play: {current.get('student_play_store_url')}\n"
            f"• App Store: {current.get('student_app_store_url')}\n\n"
            f"👨‍🏫 <b>Diamond Teachers App</b>:\n"
            f"• Minimal versiya: <code>{current.get('min_teacher_version')}</code>\n"
            f"• Minimal build: <code>{current.get('min_teacher_build')}</code>\n"
            f"• Google Play: {current.get('teacher_play_store_url')}\n"
            f"• App Store: {current.get('teacher_app_store_url')}\n\n"
            "💡 <i>Versiyani o'zgartirish uchun masalan: 'Student app versiyasini 1.2.0 build 15 qil' deb yozishingiz mumkin.</i>"
        )


def _count_student_flow_online(student_state_map: dict | None) -> int | None:
    if not student_state_map:
        return None
    n = 0
    for st in student_state_map.values():
        if not isinstance(st, dict):
            continue
        if st.get("step"):
            n += 1
    return n


def try_diamondvoy_bot_info(
    query: str,
    *,
    user: dict | None,
    telegram_user_id: int,
    lang: str,
    scope: Literal["admin_full", "student_personal", "student_limited"],
    student_state_map: dict | None = None,
) -> str | None:
    """
    If the query asks for bot/system statistics, return an HTML snippet.
    Otherwise return None (caller should use Gemini).

    - admin_full: global stats (admin bot or equivalent).
    - student_limited: ordinary student — personal data only; global aggregates denied.
    - student_personal: same restrictions as student_limited here; reserved for personal-only phrasing.
    """
    q = (query or "").strip()
    if len(q) < 3 or not _STATS_ANY.search(q):
        return None

    is_admin_context = scope == "admin_full" or telegram_user_id in ALL_ADMIN_IDS
    wants_global = True if scope == "admin_full" else bool(_GLOBAL_HINTS.search(q))
    wants_personal = bool(_PERSONAL_HINTS.search(q))

    if not is_admin_context:
        if wants_global and not wants_personal:
            return t(lang, "diamondvoy_stats_global_denied")
        # Personal / vague stats → personal snapshot
        if not user or not user.get("id"):
            return t(lang, "diamondvoy_stats_need_registration")
        uid = int(user["id"])
        total = float(get_dcoins(uid))
        dpoints_total = float(get_dpoints(uid))
        lines = [
            t(lang, "diamondvoy_stats_personal_title"),
            t(lang, "diamondvoy_stats_personal_dcoin_total", total=f"{total:.1f}"),
            t(lang, "diamondvoy_stats_personal_dpoints_total", total=f"{dpoints_total:.1f}"),
        ]
        return "\n".join(lines)

    users = get_all_users()
    n_accounts = len(users)
    n_students = len([u for u in users if u.get("login_type") in (1, 2)])
    daily_stock = count_available_daily_tests_global()
    online = _count_student_flow_online(student_state_map)
    online_line = (
        t(lang, "diamondvoy_stats_global_online", n=online)
        if online is not None
        else t(lang, "diamondvoy_stats_global_online_na")
    )
    return "\n".join(
        [
            t(lang, "diamondvoy_stats_global_title"),
            t(lang, "diamondvoy_stats_global_users", n=n_accounts),
            t(lang, "diamondvoy_stats_global_students", n=n_students),
            t(lang, "diamondvoy_stats_global_daily_stock", n=daily_stock),
            online_line,
        ]
    )


def default_subjects_for_diamondvoy(user: dict | None) -> list[str]:
    if user and user.get("id") is not None:
        from db import get_user_subjects

        subs = get_user_subjects(user["id"])
        if subs:
            return [str(s).strip().title() for s in subs]
        sub = (user.get("subject") or "English").strip().title()
        return [sub]
    return [str(s).strip().title() for s in SUBJECTS]


async def diamondvoy_analyze_student_file(
    file_bytes: bytes,
    file_name: str,
    lang: str = "uz"
) -> list[dict]:
    """
    Extracts student information (first_name, last_name, phone, parent_phone) 
    from a given .docx or .xlsx file using AI.
    Returns a list of dictionaries with keys:
      - first_name: str
      - last_name: str
      - phone: str (student phone)
      - parent_phone: str
    """
    import io
    import json
    
    extracted_text = ""
    lower_name = file_name.lower()
    
    # 1. Extract raw text from file
    try:
        if lower_name.endswith('.xlsx'):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            sheet = wb.active
            rows = []
            # Extract up to 200 rows to avoid token overflow
            for row in sheet.iter_rows(min_row=1, max_row=200, values_only=True):
                row_vals = [str(val).strip() for val in row if val is not None and str(val).strip()]
                if row_vals:
                    rows.append(" | ".join(row_vals))
            extracted_text = "\n".join(rows)
            
        elif lower_name.endswith('.docx'):
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = []
            
            # Extract text from tables first
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip().replace('\n', ' '))
                    if row_data:
                        paragraphs.append(" | ".join(row_data))
                        
            # Then extract text from paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())
                    
            extracted_text = "\n".join(paragraphs[:300]) # Limit to 300 paragraphs
        else:
            return []
            
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"Error extracting text from {file_name}: {e}")
        return []

    if not extracted_text.strip():
        return []

    # 2. Use AI to parse the text into structured JSON
    sys_prompt = (
        "Siz hujjatdan o'quvchi ma'lumotlarini ajratib oluvchi yordamchisiz. "
        "Matndan o'quvchilar ro'yxatini toping va JSON formatida qaytaring. "
        "Faqat quyidagi maydonlarga ega JSON massivini (array) qaytaring: "
        "[{'first_name': 'Ism', 'last_name': 'Familya', 'phone': 'O`quvchi teli', 'parent_phone': 'Ota-ona teli'}]. "
        "Agar telefon raqam topilmasa, bo'sh string qoldiring. "
        "Har bir obyektda ushbu 4 ta maydon albatta bo'lishi shart. "
        "Hech qanday qo'shimcha izohsiz faqat JSON qaytaring."
    )
    
    prompt = f"Hujjat matni:\n\n{extracted_text[:15000]}"
    
    try:
        import aiohttp
        from ai_generator import _xai_generate_text
        async with aiohttp.ClientSession() as session:
            ai_response = await _xai_generate_text(
                prompt,
                session=session,
                system_content=sys_prompt,
                temperature=0.1
            )
            
        # Clean up the response to extract JSON
        ai_response = ai_response.strip()
        if ai_response.startswith("```json"):
            ai_response = ai_response[7:]
        if ai_response.startswith("```"):
            ai_response = ai_response[3:]
        if ai_response.endswith("```"):
            ai_response = ai_response[:-3]
            
        ai_response = ai_response.strip()
        
        parsed_data = json.loads(ai_response)
        if isinstance(parsed_data, list):
            # Clean up phones to standard format
            for item in parsed_data:
                for key in ['phone', 'parent_phone']:
                    if item.get(key):
                        # Keep only digits and '+'
                        import re
                        cleaned = re.sub(r'[^\d+]', '', str(item[key]))
                        if not cleaned.startswith('+') and len(cleaned) >= 9:
                            cleaned = '+' + cleaned if len(cleaned) > 9 else '+998' + cleaned
                        item[key] = cleaned
            return parsed_data
            
        return []
        
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"AI parsing error for {file_name}: {e}")
        return []
